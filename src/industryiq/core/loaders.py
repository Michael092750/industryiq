"""Document loaders: turn a file on disk into plain text.

Each format has its own small, directly testable function (:func:`load_text`,
:func:`load_pdf`). :func:`load` is a dispatcher that picks the right one based
on the file extension, so callers don't need to care about the format.
"""

import logging
import re
from collections import Counter
from collections.abc import Callable
from pathlib import Path
from typing import Any

import docx
import pypdf

from industryiq.config import get_settings

logger = logging.getLogger(__name__)

# For dedup only: fold to lowercase alphanumerics so "17 .1" and a table cell "| 17.1 |"
# compare equal regardless of the parser's spacing/punctuation.
_NON_ALNUM = re.compile(r"[^a-z0-9]+")


def _dedup_key(text: str) -> str:
    return _NON_ALNUM.sub("", text.lower())


def _letters(text: str) -> int:
    return sum(ch.isalpha() for ch in text)


def load_text(path: str | Path) -> str:
    """Read a plain-text ``.txt`` file and return its contents (UTF-8).

    Raises:
        FileNotFoundError: If ``path`` does not point to an existing file.
        ValueError: If the file is not a ``.txt`` file.
    """
    p = Path(path)
    if not p.is_file():
        raise FileNotFoundError(f"No such file: {p}")
    if p.suffix.lower() != ".txt":
        raise ValueError(f"load_text expects a .txt file, got {p.suffix!r}")
    return p.read_text(encoding="utf-8")


def load_pdf_pages(path: str | Path) -> list[str]:
    """Extract a ``.pdf`` file's text page by page (one string per page).

    Which engine does the extraction is set by ``PDF_PARSER``:

    * ``"docling"`` (default) -- layout-aware parsing that emits Markdown with
      correct reading order and headings, which chunks/retrieves far better on
      multi-column report PDFs. Slower (seconds/page) and needs the optional
      ``docling`` extra (``pip install 'industryiq[docling]'``). If Docling fails
      for any reason -- not installed, or a PDF it can't convert -- this falls
      back to pypdf so a long offline ingest is never halted by one bad file.
    * ``"pypdf"`` -- fast, pure-Python text extraction with no fallback. Fine for
      clean single-column PDFs; weak on multi-column layouts.

    Either way the result is one string per page, so page-number citations work.

    Raises:
        FileNotFoundError: If ``path`` does not point to an existing file.
        ValueError: If the file is not a ``.pdf`` file.
    """
    p = Path(path)
    if not p.is_file():
        raise FileNotFoundError(f"No such file: {p}")
    if p.suffix.lower() != ".pdf":
        raise ValueError(f"load_pdf expects a .pdf file, got {p.suffix!r}")
    if get_settings().pdf_parser == "docling":
        try:
            return _load_pdf_pages_docling(p)
        except Exception as exc:  # noqa: BLE001 -- any Docling failure falls back to pypdf
            logger.warning("Docling failed on %s (%s); falling back to pypdf.", p.name, exc)
            return _load_pdf_pages_pypdf(p)
    return _load_pdf_pages_pypdf(p)


def _load_pdf_pages_pypdf(p: Path) -> list[str]:
    """Per-page plain text via pypdf (the default engine)."""
    reader = pypdf.PdfReader(str(p))
    return [page.extract_text() for page in reader.pages]


def _recover_dropped_pages(
    docling_pages: list[str],
    pypdf_pages: list[str],
    *,
    min_letters: int = 10,
    header_frac: float = 0.25,
) -> list[str]:
    """Per page, the pypdf lines whose text Docling dropped -- one block per page.

    Docling routes charts/figures/boxed callouts to ``<!-- image -->`` and discards the
    text inside them, so figure-attached prose (chart footnotes, source notes,
    average-line labels, timeline entries) is lost even though the PDF's text layer has
    it -- and the figure-VLM pass only recovers a chart's *data table*, not its footnotes.
    pypdf reads the text layer flat, so a line pypdf has that Docling doesn't is exactly
    that dropped content. Returned blocks are appended to each page's Markdown, giving a
    lossless completeness net without disturbing Docling's clean reading order.

    Dedup is whitespace/punctuation-insensitive (:func:`_dedup_key`), so text Docling kept
    -- even reflowed into a table cell or respaced -- is not re-added. Lines with fewer
    than ``min_letters`` letters (numeric table rows, axis-tick soup) are skipped: those
    are the figure-VLM's job, not prose recovery. Running headers/footers (a line repeated
    on >= ``header_frac`` of pages) are dropped so a page banner isn't re-added everywhere.
    """
    freq: Counter[str] = Counter()
    per_page_lines: list[list[str]] = []
    for text in pypdf_pages:
        lines = [" ".join(ln.split()) for ln in text.splitlines() if ln.strip()]
        per_page_lines.append(lines)
        for key in {_dedup_key(ln) for ln in lines if _letters(ln) >= min_letters}:
            freq[key] += 1
    n_pages = len(pypdf_pages)
    # A running header/footer repeats on *many* pages: require both an absolute floor
    # (>= 3 pages, so a one-off footnote on a small doc isn't mistaken for a banner) and
    # a fraction of pages (so the rule scales to long documents).
    repeated = {k for k, c in freq.items() if c >= 3 and c / n_pages >= header_frac}

    recovered: list[str] = []
    for docling_md, lines in zip(docling_pages, per_page_lines, strict=False):
        haystack = _dedup_key(docling_md)
        kept: list[str] = []
        seen: set[str] = set()
        for ln in lines:
            key = _dedup_key(ln)
            if _letters(ln) < min_letters or not key or key in seen or key in repeated:
                continue
            if key in haystack:
                continue
            seen.add(key)
            kept.append(ln)
        recovered.append("\n".join(kept))
    return recovered


def _apply_hybrid_recovery(p: Path, pages: list[str]) -> list[str]:
    """Append the text Docling dropped (per :func:`_recover_dropped_pages`) to ``pages``.

    Aligns pypdf pages to Docling pages 1:1 when the counts match (so recovered text keeps
    its page number for citations); on a count mismatch it recovers against the whole
    document and appends to the last page, so content is never lost even if attribution is
    imperfect. Best-effort: any pypdf failure logs and returns ``pages`` unchanged -- the
    recovery net must never take down a parse that already succeeded.
    """
    try:
        pypdf_pages = _load_pdf_pages_pypdf(p)
    except Exception as exc:  # noqa: BLE001 -- recovery is best-effort; keep the docling parse
        logger.warning("Hybrid recovery skipped for %s (pypdf failed: %s).", p.name, exc)
        return pages
    if len(pypdf_pages) == len(pages):
        for i, extra in enumerate(_recover_dropped_pages(pages, pypdf_pages)):
            if extra:
                pages[i] = f"{pages[i]}\n\n{extra}"
        return pages
    # Page counts disagree: recover against the whole doc, append once, lose nothing.
    whole = _recover_dropped_pages(["\n".join(pages)], ["\n".join(pypdf_pages)])
    if whole and whole[0]:
        pages[-1] = f"{pages[-1]}\n\n{whole[0]}"
    return pages


# Building a Docling converter loads ML models, so build it once and reuse it.
_docling_converter: Any = None


def _patch_rapidocr_scale(scale: int) -> None:
    """Lower the resolution at which Docling renders page regions for OCR.

    Docling's RapidOCR stage hardcodes ``self.scale = 3`` (216 DPI, then x1.5 =
    324 DPI) and exposes no option to change it, so those high-res renders pile up
    and OOM/SIGSEGV the process on large reports. We patch the model class once to
    set a lower scale. Best-effort: if Docling's internals have shifted, we log and
    leave the default rather than break ingestion.
    """
    try:  # pragma: no cover - needs the heavy extra
        from docling.models.stages.ocr import rapid_ocr_model

        if getattr(rapid_ocr_model.RapidOcrModel, "_iiq_scale_patched", False):
            return
        _orig_init = rapid_ocr_model.RapidOcrModel.__init__

        def _scaled_init(self: Any, *args: Any, **kwargs: Any) -> None:
            _orig_init(self, *args, **kwargs)
            self.scale = scale

        rapid_ocr_model.RapidOcrModel.__init__ = _scaled_init
        rapid_ocr_model.RapidOcrModel._iiq_scale_patched = True
    except Exception as exc:  # noqa: BLE001 -- Docling internals can shift across versions
        logger.warning("Could not lower RapidOCR render scale (%s); using its default.", exc)


def _soften_torch_compile() -> None:
    """Let ``torch.compile`` fall back to eager instead of failing the pipeline.

    Docling's chart-extraction and picture-description vision models ``torch.compile``
    under the hood. On a machine with no C++ compiler -- Windows without MSVC's
    ``cl.exe`` -- TorchInductor raises ``Compiler: cl is not found`` and takes down
    the whole ``StandardPdfPipeline`` (which our loader then falls back to pypdf for,
    silently losing the layout parse). ``suppress_errors`` makes a failed compile
    fall back to eager execution: slower, but it runs and produces the figure data.
    Best-effort -- if Torch's internals shift, we log and continue.
    """
    try:  # pragma: no cover - only exercised with the heavy vision extras enabled
        import torch._dynamo

        torch._dynamo.config.suppress_errors = True
    except Exception as exc:  # noqa: BLE001 -- torch internals can shift across versions
        logger.warning(
            "Could not soften torch.compile (%s); Docling figure enrichment may fail "
            "on machines without a C++ compiler.",
            exc,
        )


def _get_docling_converter() -> Any:
    """Return a cached Docling ``DocumentConverter``, built on first use.

    OCR is on by default (``DOCLING_OCR``). RapidOCR's detection step is forced to
    limit_type=max so a large embedded bitmap is downscaled (to RapidOCR's internal
    2000px ceiling) before inference -- its default (limit_type=min) only upscales,
    so a full-size chart bitmap grows the ONNX tensor until it OOMs
    (``std::bad_alloc``). Raises a pointed error if the optional ``docling`` extra
    isn't installed, rather than the bare ``ModuleNotFoundError`` the import gives.
    """
    global _docling_converter
    if _docling_converter is None:
        try:
            from docling.datamodel.base_models import InputFormat
            from docling.datamodel.pipeline_options import PdfPipelineOptions, RapidOcrOptions
            from docling.datamodel.settings import settings as docling_settings
            from docling.document_converter import DocumentConverter, PdfFormatOption
        except ModuleNotFoundError as exc:
            raise ModuleNotFoundError(
                "PDF_PARSER='docling' needs the optional 'docling' dependency; "
                "install it with:  pip install 'industryiq[docling]'"
            ) from exc
        settings = get_settings()  # pragma: no cover - needs the heavy extra
        # Serialize page rasterization to cap peak memory; the default (4 pages at
        # once) can OOM the whole page on large media, dropping its text too.
        docling_settings.perf.page_batch_size = settings.docling_page_batch_size
        if settings.docling_ocr:
            _patch_rapidocr_scale(settings.docling_ocr_scale)
        pipeline_options = PdfPipelineOptions()
        pipeline_options.do_ocr = settings.docling_ocr
        # Force the detection step to downscale large bitmaps; RapidOCR's default
        # (limit_type=min) never shrinks them, so a full-size chart bitmap OOMs.
        pipeline_options.ocr_options = RapidOcrOptions(rapidocr_params={"Det.limit_type": "max"})
        # Figure ingestion (opt-in; each runs a vision model per detected picture, so
        # both are heavy and off by default). Chart extraction turns a chart the
        # layout model tagged as a picture into a CSV of its values -- recovering the
        # numeric facts the plain layout export drops as an <!-- image --> placeholder.
        # Picture description captions the remaining (non-chart) figures. Their vision
        # models torch.compile, so soften that first (no MSVC on this box -> eager).
        if settings.docling_chart_extraction or settings.docling_picture_description:
            _soften_torch_compile()
        pipeline_options.do_chart_extraction = settings.docling_chart_extraction
        pipeline_options.do_picture_description = settings.docling_picture_description
        # FIGURE_VLM transcribes figures in a separate off-box pass (see figure_vlm.py),
        # which needs the cropped picture images. images_scale keeps the crops modest so
        # generating them doesn't blow the memory budget this pipeline already guards.
        if settings.figure_vlm != "off":
            pipeline_options.generate_picture_images = True
            pipeline_options.images_scale = 2
        _docling_converter = DocumentConverter(
            format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)}
        )
    return _docling_converter


def _load_pdf_pages_docling(p: Path) -> list[str]:
    """Per-page Markdown via Docling (layout-aware, opt-in engine).

    Falls back to a single whole-document element if Docling reports no pages. When
    ``FIGURE_VLM`` is set, a separate pass (:mod:`industryiq.core.figure_vlm`) transcribes
    each detected figure with a vision model and splices the result into the page Markdown
    in place of its ``<!-- image -->`` placeholder -- recovering chart/figure content the
    layout export otherwise drops. The pass runs *after* conversion, so a VLM failure loses
    one figure, never the document's parse.
    """
    settings = get_settings()
    doc: Any = _get_docling_converter().convert(str(p)).document

    inject = None
    figures_by_page: dict[int, list[str]] = {}
    if settings.figure_vlm != "off":
        from industryiq.core.figure_vlm import (
            annotate_document_figures,
            build_annotator,
            inject_figures,
        )

        inject = inject_figures
        figures_by_page = annotate_document_figures(
            doc,
            build_annotator(settings),
            min_pixels=settings.figure_vlm_min_pixels,
            max_figures=settings.figure_vlm_max_figures,
        )

    page_count = len(doc.pages)
    if page_count == 0:
        md = doc.export_to_markdown()
        if inject and figures_by_page:
            md = inject(md, [t for texts in figures_by_page.values() for t in texts])
        return _apply_hybrid_recovery(p, [md]) if settings.pdf_hybrid_recovery else [md]

    pages: list[str] = []
    for n in range(1, page_count + 1):
        md = doc.export_to_markdown(page_no=n)
        if inject and n in figures_by_page:
            md = inject(md, figures_by_page[n])
        pages.append(md)
    # Recover the text Docling dropped into <!-- image --> picture regions (figure
    # footnotes/annotations/callouts) from pypdf's flat text layer. Runs last, so it
    # sees the figure-VLM injections and won't re-add anything they already restored.
    return _apply_hybrid_recovery(p, pages) if settings.pdf_hybrid_recovery else pages


def load_pdf(path: str | Path) -> str:
    """Extract text from a ``.pdf`` file, joining pages with newlines."""
    return "\n".join(load_pdf_pages(path))


def load_docx(path: str | Path) -> str:
    """Extract text from a ``.docx`` file, joining paragraphs with newlines.

    Raises:
        FileNotFoundError: If ``path`` does not point to an existing file.
        ValueError: If the file is not a ``.docx`` file.
    """
    p = Path(path)
    if not p.is_file():
        raise FileNotFoundError(f"No such file: {p}")
    if p.suffix.lower() != ".docx":
        raise ValueError(f"load_docx expects a .docx file, got {p.suffix!r}")
    document = docx.Document(str(p))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


_LOADERS: dict[str, Callable[[str | Path], str]] = {
    ".txt": load_text,
    ".pdf": load_pdf,
    ".docx": load_docx,
}

SUPPORTED_EXTENSIONS = frozenset(_LOADERS)


def _to_utf8_safe(text: str) -> str:
    """Drop characters that cannot be encoded as UTF-8.

    PDF extraction can emit lone surrogate code points (from broken font maps)
    that are valid in a Python ``str`` but not encodable to UTF-8. Left in, they
    crash every UTF-8 consumer downstream -- the embedding tokenizer, JSON
    payloads to Bedrock, and Postgres text columns. Stripping them here keeps
    each loader's output safe to embed and store.
    """
    return text.encode("utf-8", "ignore").decode("utf-8")


def load(path: str | Path) -> str:
    """Load any supported file by dispatching on its extension.

    The returned text is guaranteed UTF-8 encodable (see :func:`_to_utf8_safe`).

    Raises:
        FileNotFoundError: If ``path`` does not point to an existing file.
        ValueError: If the file's extension is not supported.
    """
    p = Path(path)
    loader = _LOADERS.get(p.suffix.lower())
    if loader is None:
        raise ValueError(
            f"Unsupported file type {p.suffix!r}; supported: {sorted(SUPPORTED_EXTENSIONS)}"
        )
    return _to_utf8_safe(loader(path))


def load_pages(path: str | Path) -> list[str]:
    """Load a file as a list of page texts (each UTF-8 safe), for page attribution.

    PDFs return one element per page; other formats have no real pagination, so
    they return the whole document as a single element.
    """
    p = Path(path)
    if p.suffix.lower() == ".pdf":
        return [_to_utf8_safe(page) for page in load_pdf_pages(p)]
    return [load(p)]


def load_title(path: str | Path) -> str | None:
    """The document's embedded title, or ``None`` if it has none.

    Reads PDF/DOCX document metadata; other formats have no title. Callers
    typically fall back to the file name. Never raises -- a missing or unreadable
    title is just ``None``.
    """
    p = Path(path)
    ext = p.suffix.lower()
    try:
        if ext == ".pdf":
            info = pypdf.PdfReader(str(p)).metadata
            title = info.title if info else None
        elif ext == ".docx":
            title = docx.Document(str(p)).core_properties.title
        else:
            return None
    except Exception:  # noqa: BLE001 -- best-effort metadata read; absence is fine
        return None
    title = (title or "").strip()
    return title or None
